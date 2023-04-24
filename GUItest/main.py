# 这是一个示例 Python 脚本。


# 按 Shift+F10 执行或将其替换为您的代码。
# 按 双击 Shift 在所有地方搜索类、文件、工具窗口、操作和设置。

import tkinter as tk
import os
import win32com.client
from tkinter import *
from tkinter import filedialog
from tkinter import messagebox, NW
from PIL import Image, ImageTk
from osgeo import gdal
import pymongo
import docx


class Main():
    def __init__(self):  # 创建新窗口，初始化窗口
        self.mwt = tk.Tk()
        self.mwt.title('test')
        self.mwt.geometry('800x1000')
        self.vax = tk.IntVar()
        self.vay = tk.IntVar()
        self.errorinfo = errorinfo()
        self.setmenu()  # 设置菜单栏

        self.mwt.mainloop()  # 显示窗口

    # 设置总菜单
    def setmenu(self):  # 设置菜单栏
        self.mainmenu = tk.Menu(self.mwt)

        self.loadmenu()
        self.errorsavemenu()
        # self.errorsearchmenu()


        self.mainmenu.add_cascade(label='打开文件', menu=self.loadmenu)
        self.mainmenu.add_cascade(label='错误标记', menu=self.errorsavemenu)
        self.mainmenu.add_command(label='错误查询', command=self.errorsearchmenu)
        self.mainmenu.add_command(label='help', command=self.help1menu)

        self.mwt.config(menu=self.mainmenu)

    # ----打开文件菜单---------------------------------------------------------------------------------------------------------
    def loadmenu(self):  # load下拉菜单
        self.loadmenu = tk.Menu(self.mainmenu, tearoff=0)
        self.setframe()  # 设置页面布局

        self.loadmenu.add_command(label='image', command=self.open_image)
        self.loadmenu.add_command(label='docx', command=self.open_docx)
        self.loadmenu.add_command(label='txt', command=self.open_txt)

    # 打开图片功能
    def open_image(self):  # 只能打开tif类型文件
        imagefilename = filedialog.askopenfilename()
        # 打开数据集
        dataset = gdal.Open(imagefilename)
        # 获取数据集的宽度和高度
        self.imgwidth = dataset.RasterXSize
        self.imgheight = dataset.RasterYSize


        # 将栅格数据转换为PIL图像对象
        bands = dataset.RasterCount
        if bands == 1:  # 灰度图像
            img = Image.fromarray(dataset.ReadAsArray())
        else:  # 彩色图像
            red = dataset.GetRasterBand(1).ReadAsArray()
            green = dataset.GetRasterBand(2).ReadAsArray()
            blue = dataset.GetRasterBand(3).ReadAsArray()
            img = Image.merge("RGB", (Image.fromarray(red), Image.fromarray(green), Image.fromarray(blue)))
        self.mwt.geometry('{}x{}+20+20'.format(img.width * 2, img.height))
        self.mwt.update()

        # 创建 PhotoImage 对象
        photo_img = ImageTk.PhotoImage(img)
        self.lable_img = Label(self.lefrm, image=photo_img)
        self.lable_img.photo = photo_img
        self.lable_img.place(x=0,y=0)
        if not self.lable_img:
            messagebox.showinfo('提示','打开图片失败！请重新选择图片格式（支持tif格式）')

        def show_coordinates(event):
            x, y = event.x, event.y
            self.vax.set(x)
            self.vay.set(y)

        self.lable_img.bind("<Button-1>", show_coordinates)

        self.mwt.mainloop()

    # 打开docx文件
    def open_docx(self):

        docfilename = filedialog.askopenfilename()
        # 获取文件扩展名
        ext = os.path.splitext(docfilename)[1]
        # 如果是docx文件
        if ext == '.docx':
            document = docx.Document(docfilename)
        # 如果是doc文件
        elif ext == '.doc':
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(docfilename)
            doc.SaveAs(os.path.splitext(docfilename)[0] + '.docx', FileFormat=12)
            doc.Close()
            document = docx.Document(os.path.splitext(docfilename)[0] + '.docx')
        else:
            raise ValueError('Unsupported file format')

        # 加载文本文件
        text = ''
        for paragraph in document.paragraphs:
            text += paragraph.text + '\n'

        # 在窗口中显示文本内容
        text_widget = tk.Text(self.lefrm, height=100, width=100, font=('Arial', 12))
        text_widget.insert('1.0', text)
        text_widget.place(anchor=NW)

        def show_coordinates(event):
            x, y = event.x, event.y
            self.vax.set(x)
            self.vay.set(y)
        text_widget.bind("<Button-1>", show_coordinates)

    # 打开txt文件
    def open_txt(self):
        file_path = filedialog.askopenfilename()
        with open(file_path, 'r') as file:
            text = file.read()
            text_widget = tk.Text(self.lefrm)
            text_widget.delete('1.0', 'end')
            text_widget.insert('end', text)
            text_widget.pack()

    # ----错误标注菜单---------------------------------------------------------------------------------------------------------
    def errorsavemenu(self):  # 错误标注下拉菜单
        self.setframe()  # 设置页面布局
        self.errorsavemenu = tk.Menu(self.mainmenu, tearoff=0)
        self.errorsavemenu.add_command(label='图片', command=self.error_img)
        self.errorsavemenu.add_command(label='docx文本', command=self.error_file)


    def error_img(self):
        self.error_input()
        self.open_image()


    def error_file(self):
        self.error_input()
        self.open_docx()

    # 输入错误信息并保存
    def error_input(self):
        self.mwt.update()
        # 错误信息输入
        self.eL1 = Label(self.rifrm, text='错误坐标:', height=2)
        self.eL1.grid(row=3, column=0)
        self.L2 = tk.Label(self.rifrm, textvariable=self.vax, width=20, height=2)
        self.L2.grid(row=4, column=0)
        # self.L2.pack()
        self.L3 = tk.Label(self.rifrm, textvariable=self.vay, width=20, height=2)
        self.L3.grid(row=4, column=1)
        # self.L3.pack()
        self.eL2 = Label(self.rifrm, text='错误类型:', height=2)
        self.eL2.grid(row=5, column=0)
        self.eEn1 = Entry(self.rifrm)
        self.eEn1.grid(row=6, rowspan=2, column=0, columnspan=2)
        self.eL3 = Label(self.rifrm, text='错误描述:', height=2)
        self.eL3.grid(row=8, column=0)
        self.eEn2 = Entry(self.rifrm, width=25)
        self.eEn2.grid(row=10, rowspan=2, column=0, columnspan=2)
        # 如何指定文件名称
        self.eL4 = Label(self.rifrm, text='文件名：', height=2)
        self.eL4.grid(row=12, column=0)
        self.eEn3 = Entry(self.rifrm, width=25)
        self.eEn3.grid(row=13, rowspan=2, column=0, columnspan=2)

        # 保存错误信息
        def errorinfo_save():
            if self.errorinfo.index==0:
                filename = self.eEn3.get()
                self.errorinfo.filename = filename + '.txt'
                self.eEn3.grid_forget()
            else:
                self.eEn3.destroy()
            # print(self.erroinfo.filename)
            type = self.eEn1.get()
            describe = self.eEn2.get()
            if not describe:
                describe=' '
            self.errorinfo.errx.insert(self.errorinfo.index, int(self.vax.get()))
            self.errorinfo.erry.insert(self.errorinfo.index, int(self.vay.get()))
            if not type:
                self.errorinfo.errtype.append(' ')
            else:
                self.errorinfo.errtype.insert(self.errorinfo.index, type)
            if not type:
                self.errorinfo.errcon.append(' ')
            else:
                self.errorinfo.errcon.insert(self.errorinfo.index, describe)
            self.errorinfo.errindex.insert(self.errorinfo.index, self.errorinfo.index)
            with open(self.errorinfo.filename, 'a') as file:
                file.write(str(self.errorinfo.index) + " ")
                file.write(str(self.errorinfo.errx[self.errorinfo.index]) + ' ' + str(
                    self.errorinfo.erry[self.errorinfo.index]) + ' ')
                file.write(str(self.errorinfo.errtype[self.errorinfo.index]) + ' ' + str(
                    self.errorinfo.errcon[self.errorinfo.index]) + '\n')
            tk.messagebox.showinfo("提示", "保存成功！")
            file.close()
            self.L4 = tk.Label(self.rifrm, text='共有' + str(self.errorinfo.index + 1) + '条错误信息', width=20,
                               height=2)
            self.L4.grid(row=17, column=0, columnspan=2)
            self.eL4.destroy()
            self.errorinfo.index += 1

        # 清除错误信息
        def errorinfo_clear():
            self.eEn1.delete(0, tk.END)
            self.eEn2.delete(0, tk.END)

        self.eSaveB1 = Button(self.rifrm, text='保存', command=errorinfo_save)
        self.eSaveB1.grid(row=15, rowspan=2, column=0)
        self.eSaveB2 = Button(self.rifrm, text='放弃', command=errorinfo_clear)
        self.eSaveB2.grid(row=15, rowspan=2, column=1)


    # ----错误查询菜单--------------------------------------------------------------------------------------------------------
    def errorsearchmenu(self):
        # self.errorsearchmenu=tk.Menu(self.mw)
        # self.errorsearchmenu.add_command(label='打开错误文件',command=self.error_search)
            self.sc = tk.Toplevel()
            self.sc.geometry('500x500')
            filename = filedialog.askopenfilename()
            Lfile = Label(self.sc, text='文件路径：' + str(filename), width=50, height=3)
            Lfile.grid(row=0, column=0, columnspan=5)
            # self.errorinfo=errorinfo()
            self.errorinfo.readfile(filename)
            L1 = Label(self.sc, text='请输入需要查找的关键字类型：')
            L1.grid(row=1, column=0, columnspan=3)
            options = ['坐标', '错误类型', '错误描述']
            # 创建下拉选项框
            self.var = tk.StringVar(self.sc)
            self.var.set(options[0])
            option_menu = tk.OptionMenu(self.sc, self.var, *options)
            option_menu.grid(row=1, column=4, columnspan=2)
            Lx = Label(self.sc, text='x:', height=3)
            Lx.grid(row=3, column=0)
            self.Ex = Entry(self.sc)
            self.Ex.grid(row=3, column=1, columnspan=3)
            Ly = Label(self.sc, text='y:')
            Ly.grid(row=3, column=4)
            self.Ey = Entry(self.sc)
            self.Ey.grid(row=3, column=5)
            Lt = Label(self.sc, text='错误类型:', height=3)
            Lt.grid(row=4, column=0)
            self.Et = Entry(self.sc)
            self.Et.grid(row=4, column=1, columnspan=3)
            Ld = Label(self.sc, text='错误描述:', height=3)
            Ld.grid(row=5, column=0)
            self.Ed = Entry(self.sc)
            self.Ed.grid(row=5, column=1, columnspan=3)

            Bsel = tk.Button(self.sc, text='确认', command=self.confirm)
            Bsel.grid(row=6, column=1, columnspan=3)


            self.sc.mainloop()

    # 查找错误信息
    def confirm(self):
        self.selection = str(self.var.get())
        result=errorinfo()
        if self.selection == '坐标':
            x = self.Ex.get()
            y = self.Ey.get()
            for i in range(len(self.errorinfo.errtype)):
                if x == self.errorinfo.errx[i] and y == self.errorinfo.erry[i]:
                    result.errindex.append(self.errorinfo.errindex[i])
                    result.errx.append(self.errorinfo.errx[i])
                    result.erry.append(self.errorinfo.erry[i])
                    result.errtype.append(self.errorinfo.errtype[i])
                    result.errcon.append(self.errorinfo.errcon[i])
            if not self.errorinfo.errtype:
                print('no')

        elif self.selection == '错误类型':
            type = self.Et.get()
            for i in range(len(self.errorinfo.errtype)):
                if type == self.errorinfo.errtype[i]:
                    result.errindex.append(self.errorinfo.errindex[i])
                    result.errx.append(self.errorinfo.errx[i])
                    result.erry.append(self.errorinfo.erry[i])
                    result.errtype.append(self.errorinfo.errtype[i])
                    result.errcon.append(self.errorinfo.errcon[i])
            if not result.errtype:
                print('no')
        elif self.selection == '错误描述':
            con = self.Ed.get()
            for i in range(len(self.errorinfo.errcon)):
                if con == self.errorinfo.errcon[i]:
                    result.errindex.append(self.errorinfo.errindex[i])
                    result.errx.append(self.errorinfo.errx[i])
                    result.erry.append(self.errorinfo.erry[i])
                    result.errtype.append(self.errorinfo.errtype[i])
                    result.errcon.append(self.errorinfo.errcon[i])

        self.Lresult = Label(self.sc, text='查找结果: 共' + str(len(result.errx)) + '项结果', height=3)
        self.Lresult.grid(row=7, column=0, columnspan=3)
        self.Tresult=tk.Text(self.sc, width=50, height=5)
        if not self.errorinfo.errtype:
            self.Tresult.insert('1.0', ' ')
        else:
            for i in range(len(result.errx)):
               self.Tresult.insert(tk.END, ' {} {} {} {}\n'.format(result.errx[i], result.erry[i],result.errtype[i], result.errcon[i]))
        self.Tresult.grid(row=8, column=0, columnspan=5)

        self.sc.mainloop()



    # ----帮助菜单-----------------------------------------------------------------------------------------------------------
    def help1menu(self):  # help下拉菜单
        messagebox.showinfo('帮助','您可以使用’打开‘菜单打开并查看图像或文本。\n使用’错误标注‘菜单对图像或文本进行错误标注，并生成错误文件信息。\n使用’错误查询‘菜单对错误信息进行查询。')


    # ----界面布局设置菜单-----------------------------------------------------------------------------------------------------
    def setframe(self):
        self.maiframe = tk.Frame(self.mwt).pack()
        self.setleftfrm()
        self.setrightfrm()

    # 设置左布局
    def setleftfrm(self):
        self.mwt.update()
        self.lefrm = tk.Frame(self.maiframe, height=self.mwt.winfo_height(), width=self.mwt.winfo_width() / 2)
        self.lefrm.place(anchor=NW)

    # 设置右布局
    def setrightfrm(self):
        self.mwt.update()
        width1=self.mwt.winfo_width()
        self.rifrm = tk.Frame(self.maiframe, height=self.mwt.winfo_height(), width=width1/2)
        self.rifrm.place(x=width1/2,y=0)

# ---错误文件类------------------------------------------------------------------------------------------------------------
class errorinfo():
    def __init__(self):#初始化错误信息类
        self.filename = ''
        self.errx = []
        self.erry = []
        self.errtype = []
        self.errcon = []
        self.errindex = []
        self.index = 0


    #读入错误信息文件
    def readfile(self,filename):
        with open(filename,'r') as A:
            for eachline in A:
                tmp = eachline.split()
                self.errindex.append(tmp[0])
                self.errx.append(tmp[1])
                self.erry.append(tmp[2])
                self.errtype.append(tmp[3])
                self.errcon.append(tmp[4])
                print(eachline)
        messagebox.showinfo('提示','读入成功！')


Main()





